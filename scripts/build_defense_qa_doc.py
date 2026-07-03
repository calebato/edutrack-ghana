from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Defense_QA_Answers_EduTrack_Ghana.docx")


SECTIONS = [
    (
        "Core Project Questions",
        [
            (
                "What is the title of your project?",
                "The title of our project is Personalized Learning System Using Machine Learning in Education - Adapt Teaching to Individual Student Needs. The developed system is called EduTrack Ghana, a web-based learning platform for Ghanaian Junior High School students.",
            ),
            (
                "What problem does EduTrack Ghana solve?",
                "EduTrack Ghana solves the problem of one-size-fits-all teaching, where all students receive the same learning materials even though they learn at different speeds and have different weaknesses. It helps students learn through quizzes, progress tracking, recommendations, badges, and teacher support.",
            ),
            (
                "Why did you choose Ghanaian Junior High Schools as your focus?",
                "We chose Ghanaian Junior High Schools because this level is important for preparing students for BECE and future academic work. Many JHS learners need more personalized support, and teachers also need better tools to identify students who are struggling early.",
            ),
            (
                "What is wrong with the traditional one-size-fits-all teaching approach?",
                "The one-size-fits-all approach assumes that all students understand lessons at the same pace and in the same way. In reality, some students need more practice, some need easier explanations, and others are ready for more challenging work, so the approach can increase learning gaps.",
            ),
            (
                "What is the main aim of your project?",
                "The main aim is to design and develop a machine learning-based Personalized Learning System that adapts teaching and learning support to the individual needs of Ghanaian Junior High School students.",
            ),
            (
                "What are your specific objectives?",
                "The objectives are to design a learner profiling module with gamification, develop a machine learning-based recommendation engine, create a teacher analytics dashboard with performance prediction and parent reports, and design a scalable framework for real-time personalized learning.",
            ),
            (
                "Who are the main users of the system?",
                "The main users are students, teachers, and system administrators. Students learn and take quizzes, teachers monitor performance and create reports, and administrators manage accounts, subjects, topics, approvals, announcements, and logs.",
            ),
            (
                "Why is machine learning relevant to your project?",
                "Machine learning is relevant because the system uses student learning data to support personalization. It helps recommend topics, profile learners, and estimate exam readiness so teachers can make better support decisions.",
            ),
            (
                "What makes EduTrack different from existing systems like DreamBox or ClickUp?",
                "EduTrack is designed specifically for Ghanaian JHS education and combines learning content, quizzes, teacher analytics, parent reporting, gamification, and machine learning in one local PHP/MySQL system. It is also lightweight enough for a school-based XAMPP environment.",
            ),
            (
                "Why did you choose the name EduTrack Ghana?",
                "The name combines education and tracking. It shows that the system helps track student learning progress, quiz performance, recommendations, and teacher reports within the Ghanaian education context.",
            ),
        ],
    ),
    (
        "System And Features",
        [
            (
                "What are the main modules of your system?",
                "The main modules are authentication, student dashboard, teacher dashboard, administrator dashboard, quiz and topic management, progress tracking, reporting, announcements, violation reports, and machine learning services.",
            ),
            (
                "What can a student do in the system?",
                "A student can register or log in, view subjects and topics, take quizzes, track progress, earn badges, view leaderboards, receive recommendations, set learning goals, use accessibility tools, and report violations.",
            ),
            (
                "What can a teacher do in the system?",
                "A teacher can monitor students, create topics and quizzes, view analytics, send announcements, check individual student details, generate progress reports, and email reports to parents or guardians.",
            ),
            (
                "What can an administrator do in the system?",
                "An administrator can manage students, teachers, subjects, topics, announcements, violations, and system logs. Admins can also approve teachers and review teacher-created curriculum content before it becomes visible to students.",
            ),
            (
                "How does the quiz system work?",
                "A student selects an available quiz, the system chooses questions, stores the selected question IDs, and then scores the answers after submission. The score is saved, progress and mastery are updated, and the student may earn points or badges.",
            ),
            (
                "How does the system track student progress?",
                "The system records completed topics, quiz attempts, scores, mastery levels, badges, streaks, and learning activity. These records are used to show progress dashboards and support recommendations.",
            ),
            (
                "How does the badge or gamification system help students?",
                "Badges, points, streaks, and leaderboards motivate students to continue learning. They make progress visible and encourage students to complete topics, pass quizzes, and maintain consistent study habits.",
            ),
            (
                "How are parent reports generated?",
                "Teachers generate reports from student profile data, quiz history, subject performance, strengths, weaknesses, badges, remarks, and available ML forecasts. The report can be printed or emailed to a parent or guardian.",
            ),
            (
                "Why did you include violation reporting?",
                "Violation reporting gives students a channel to report issues such as misconduct or unsafe behavior. It supports school accountability and allows teachers or administrators to review and respond to reported cases.",
            ),
            (
                "How does the system support accessibility?",
                "The system includes an accessibility page and speech transcription support. Audio can be sent through the transcription endpoint so learners who prefer voice input or need support can convert speech to text.",
            ),
        ],
    ),
    (
        "Machine Learning Questions",
        [
            (
                "What machine learning algorithms did you use?",
                "The system includes a Ridge Regression baseline predictor, a personal linear trend fallback, an XGBoost prediction and risk model, a TensorFlow learner profile model, a TensorFlow contextual bandit recommendation model, and Whisper for speech transcription.",
            ),
            (
                "What data does the system use to make predictions?",
                "It uses data such as quiz attempt count, average score, recent score average, score trend, pass rate, average quiz time, topic mastery, topic completion, login activity, and current learning streak.",
            ),
            (
                "How does the system recommend content to students?",
                "The recommender looks at mastery gaps, quiz scores, topic progress, difficulty fit, learning goals, and risk level. If the advanced ML service is available, it ranks topics using the contextual bandit model; otherwise PHP fallback rules provide recommendations.",
            ),
            (
                "What does the exam prediction actually predict?",
                "It predicts a student's exam readiness score, possible BECE grade band, risk level, confidence, and explanation factors. It is meant to guide study and teacher intervention, not to replace an official exam result.",
            ),
            (
                "Is the predicted score the same as a final exam result?",
                "No. The predicted score is only a planning estimate based on available learning data. Official exam results still come from actual exams, and teachers should use the prediction as support, not as the final judgment.",
            ),
            (
                "How do you handle students with little or no quiz data?",
                "The system does not force a prediction for new learners. If there is insufficient evidence, it returns an insufficient data status and encourages the learner to complete more quizzes before a forecast is shown.",
            ),
            (
                "What happens if the ML service is offline?",
                "The main PHP system continues to work. EduTrack has fallback logic in the PHP ML layer, so quizzes, progress tracking, reports, and basic recommendations can still run even when the Flask service is unavailable.",
            ),
            (
                "How accurate is your machine learning model?",
                "The model is a prototype and its accuracy depends on the amount and quality of student performance data. We designed safeguards and explainable outputs, but stronger validation will require more verified final exam results from real school use.",
            ),
            (
                "What are the limitations of your ML model?",
                "The limitations include limited real-world training data, few verified final exam labels, possible bias from small datasets, and the need for more testing across different schools. The Whisper model is pretrained and should not be claimed as fully Ghanaian-accent fine-tuned unless that training is completed and evaluated.",
            ),
            (
                "How would you improve the model in the future?",
                "We would collect more verified exam results, expand pilot testing to more schools, evaluate model accuracy carefully, improve bias checks, fine-tune speech recognition with consented Ghanaian-accent audio, and retrain model versions over time.",
            ),
        ],
    ),
    (
        "Technical Questions",
        [
            (
                "What technologies did you use and why?",
                "We used PHP for the backend, MySQL for the database, HTML5, CSS3, Bootstrap 5, and JavaScript for the frontend, XAMPP for local deployment, and Python Flask for machine learning services. These tools are accessible, lightweight, and suitable for a school prototype.",
            ),
            (
                "Why did you use PHP and MySQL?",
                "PHP and MySQL are widely used for web applications, easy to host locally with XAMPP, and suitable for CRUD operations such as accounts, quizzes, topics, reports, and logs. They also fit the project timeline and school-based deployment environment.",
            ),
            (
                "Why did you use XAMPP for deployment?",
                "XAMPP provides Apache, PHP, and MySQL in one local package, making it easy to run and demonstrate the system without needing cloud hosting. It is appropriate for development and pilot demonstration, although production deployment would need stronger server configuration.",
            ),
            (
                "How is your database structured?",
                "The database is structured around users, schools, curriculum, quizzes, learning activity, rewards, communication, audit logs, and ML outputs. Relationships connect students to schools, subjects to topics, topics to quizzes, quizzes to questions, and students to quiz attempts.",
            ),
            (
                "What are the most important database tables?",
                "Important tables include students, teachers, admins, schools, subjects, topics, quizzes, questions, quiz_attempts, topic_progress, student_learning_profiles, badges, student_badges, announcements, activity_logs, and the ML tables such as student_predictions and student_recommendations.",
            ),
            (
                "How do students, quizzes, topics, and scores relate in the database?",
                "Subjects contain topics, topics contain quizzes, and quizzes contain questions. When a student takes a quiz, the result is stored in quiz_attempts, and that score can update the student's mastery and progress for the related topic.",
            ),
            (
                "What is the purpose of the quiz_attempts table?",
                "The quiz_attempts table stores each student's quiz attempt, including the score, answers, selected question IDs, start time, completion time, and pass status. This makes scoring traceable and supports progress tracking and ML feature extraction.",
            ),
            (
                "What is the purpose of student_learning_profiles?",
                "student_learning_profiles stores each student's mastery level for specific topics. It helps the system know which topics a learner understands well and which topics need more practice or recommendation.",
            ),
            (
                "What is the purpose of the ML tables?",
                "The ML tables store model metadata, cached predictions, recommendations, learner profiles, learning goals, and verified final exam results. This keeps ML outputs traceable by model version and avoids recalculating everything every time.",
            ),
            (
                "How does the PHP application communicate with the Python/Flask ML service?",
                "PHP sends requests to the local Flask API through the ML client and bridge endpoints. The Flask service receives learner features, runs prediction, profiling, recommendation, or transcription logic, and returns JSON responses to PHP.",
            ),
        ],
    ),
    (
        "Security Questions",
        [
            (
                "How do you protect user passwords?",
                "Passwords are stored as hashes, not plain text. The system uses PHP password hashing functions when creating accounts and password_verify during login.",
            ),
            (
                "How do you prevent SQL injection?",
                "The system uses PDO prepared statements through helper functions such as dbQuery, dbRow, dbRows, and dbInsert. User input is passed as parameters instead of being directly joined into SQL strings.",
            ),
            (
                "What is CSRF protection and where is it used?",
                "CSRF protection prevents attackers from tricking a logged-in user into submitting unwanted form actions. EduTrack generates CSRF tokens and validates them before important POST actions such as login, registration, quiz submission, profile updates, and administrative changes.",
            ),
            (
                "How do you prevent students from accessing teacher/admin pages?",
                "Protected pages use role-based guards such as requireStudent, requireTeacher, and requireAdmin. These functions check the active session role and redirect or block users who do not have permission.",
            ),
            (
                "How do you protect sensitive student data?",
                "The system protects student data through login requirements, role-based access, prepared statements, output escaping, school-based teacher restrictions, and audit logs. In production, HTTPS, stronger server settings, and careful data-sharing policies would also be required.",
            ),
            (
                "Why is role-based access control important in your system?",
                "Role-based access control ensures that each user only sees and performs actions meant for their role. For example, students should not approve teachers, and teachers should not access admin system management pages.",
            ),
            (
                "What security improvements would be needed before real production deployment?",
                "Production would require HTTPS, environment-specific database credentials, least-privilege database users, restricted error output, removal or blocking of development utilities, stronger backup policies, and careful protection of real student and parent data.",
            ),
        ],
    ),
    (
        "Methodology And Design",
        [
            (
                "Why did you choose Agile Scrum?",
                "We chose Agile Scrum because it supports iterative development, continuous feedback, and gradual improvement. This allowed us to build core features, test them, receive feedback, and improve the system in sprints.",
            ),
            (
                "What are your functional requirements?",
                "Functional requirements include registration, login, learner profiling, subject and topic access, quizzes, progress tracking, recommendations, badges, teacher analytics, parent reports, announcements, admin management, and ML prediction.",
            ),
            (
                "What are your non-functional requirements?",
                "Non-functional requirements include usability, security, reliability, scalability, accessibility, fast response time, and support for low-resource school environments.",
            ),
            (
                "Explain your context diagram.",
                "The context diagram shows the system as a central platform interacting with students, teachers, administrators, parents, the database, and the ML service. It gives a high-level view of how external users and services exchange information with EduTrack.",
            ),
            (
                "Explain your ERD.",
                "The ERD shows the database relationships. For example, schools have students and teachers, subjects have topics, topics have quizzes, quizzes have questions, students make quiz attempts, and students earn badges and receive predictions or recommendations.",
            ),
            (
                "Explain your DFD.",
                "The Data Flow Diagram shows how data moves through the system. For example, a student submits quiz answers, the system stores attempts in the database, updates progress, sends features to the ML layer, and displays feedback or recommendations.",
            ),
            (
                "Explain your use case diagram.",
                "The use case diagram shows what each actor can do. Students can learn, take quizzes, and view progress; teachers can create quizzes, view analytics, and generate reports; admins can manage users, subjects, topics, announcements, and logs.",
            ),
            (
                "What testing methods did you use?",
                "We used unit testing, functional testing, usability testing, and acceptance testing. We also ran focused PHP regression checks for validation, CSRF, roles, model loading, prediction safeguards, recommendations, and metadata.",
            ),
            (
                "Why did you choose pilot implementation?",
                "Pilot implementation is suitable because the system should first be tested in a limited school environment before wider deployment. This helps identify usability issues, data issues, and technical improvements before scaling.",
            ),
        ],
    ),
    (
        "Limitations And Future Work",
        [
            (
                "What are the limitations of your project?",
                "The limitations include limited pilot scope, limited real-world ML training data, local XAMPP deployment, no long-term impact evaluation, possible device and connectivity constraints, and the need for more production security hardening.",
            ),
            (
                "Why was the system limited to one pilot school?",
                "It was limited to one pilot school to make testing manageable within the project timeline. Starting small helps us evaluate the system carefully before expanding to more schools.",
            ),
            (
                "What would you improve if given more time?",
                "We would improve mobile support, deploy online with HTTPS, collect more verified exam data, improve ML accuracy, add SMS notifications for parents, expand accessibility, and test the system in more schools.",
            ),
            (
                "How can the system be scaled to more schools?",
                "The system can be scaled by deploying it on a secure online server, adding multi-school administration, improving hosting resources, using stronger database roles, and creating onboarding processes for schools, teachers, and students.",
            ),
            (
                "Would you make it a mobile app?",
                "Yes. A mobile app would make the system easier for students and parents to access, especially where phones are more common than computers. The current responsive web design is a good foundation for later mobile support.",
            ),
            (
                "How would you improve accessibility?",
                "We would improve accessibility by adding better screen-reader support, clearer keyboard navigation, more local language support, stronger audio transcription, larger text options, and offline-friendly learning materials.",
            ),
            (
                "How would you improve parent communication?",
                "We would add SMS notifications, WhatsApp integration where appropriate, scheduled parent report delivery, and clearer parent dashboards. This would help parents follow student progress even if they do not log in often.",
            ),
            (
                "What data would you need to make the ML model better?",
                "We would need more verified final exam results, more quiz histories, topic mastery records, attendance or engagement indicators, and data from different schools. The data should be collected ethically with consent and privacy protection.",
            ),
        ],
    ),
    (
        "Hot Questions To Prepare Very Well",
        [
            (
                "If the system is local on XAMPP, how can schools use it in real life?",
                "XAMPP was used for development and demonstration. In real life, the system can be deployed on a secure school server or cloud server with HTTPS, proper database credentials, backups, and controlled access for students, teachers, and administrators.",
            ),
            (
                "How do you know the recommendations are correct?",
                "The recommendations are based on observable learning evidence such as mastery gaps, quiz performance, progress, and learning goals. They are useful study guidance, but they should be improved and validated over time using student outcomes and teacher feedback.",
            ),
            (
                "What happens when a student performs badly because of internet/device issues, not ability?",
                "That is a valid concern. The system should not treat one poor attempt as the full picture. It uses multiple attempts and trends, and teachers should interpret reports with context such as device problems, attendance, internet access, or personal challenges.",
            ),
            (
                "How do you avoid bias in the ML prediction?",
                "We reduce risk by using learning activity and performance features rather than sensitive personal characteristics. However, bias can still occur if the training data is small or unbalanced, so future work should include larger datasets, fairness checks, and teacher review.",
            ),
            (
                "Why should teachers trust the prediction?",
                "Teachers should trust it as a support tool, not as a final decision. The prediction includes confidence, risk level, model version, and explanation factors, so teachers can understand the reason and combine it with their professional judgment.",
            ),
            (
                "Which part of the project did you personally contribute to?",
                "Answer this honestly based on your actual role. A strong answer is: I contributed to system development, database setup, user interface testing, documentation, and preparing the defense materials. If asked for details, mention the exact module or task you worked on most.",
            ),
        ],
    ),
]


def set_run(run, size=None, bold=False, color=None):
    font = run.font
    font.name = "Calibri"
    if size:
        font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor.from_string(color)


def add_answer(doc, number, question, answer):
    p = doc.add_paragraph()
    p.style = "Heading 3"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{number}. {question}")
    set_run(run, 12, True, "1F4D78")

    a = doc.add_paragraph()
    a.style = "Normal"
    a.paragraph_format.left_indent = Inches(0.15)
    a.paragraph_format.space_after = Pt(6)
    run = a.add_run(answer)
    set_run(run, 11, False, "000000")


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("EduTrack Ghana Defense Questions And Answers")
    set_run(r, 22, True, "0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Prepared for project defense on June 30, 2026")
    set_run(r, 11, False, "555555")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    r = note.add_run(
        "Study tip: do not memorize every sentence word-for-word. Understand the idea, then answer naturally using your own voice."
    )
    set_run(r, 11, True, "000000")

    intro = doc.add_paragraph()
    intro.add_run(
        "This guide contains concise answers to likely panel questions based on the project report, defense slides, and EduTrack Ghana codebase. The strongest answers are the ones that explain what the system does, why it matters, how it works, and what can be improved."
    )

    doc.add_heading("Quick 2-Minute Summary", level=1)
    summary = doc.add_paragraph()
    summary.add_run(
        "EduTrack Ghana is a web-based personalized learning system for Ghanaian Junior High School students. It supports students, teachers, and administrators. Students can study topics, take quizzes, track progress, earn badges, and receive recommendations. Teachers can monitor performance, create quizzes, view analytics, and generate reports for parents. Administrators manage users, subjects, topics, approvals, announcements, violations, and logs. The system uses PHP and MySQL for the main web application, HTML/CSS/Bootstrap/JavaScript for the interface, XAMPP for local deployment, and Python Flask for optional machine learning services. The ML layer supports learner profiling, exam readiness prediction, topic recommendations, and speech transcription, while PHP fallbacks keep the platform usable if the ML service is offline."
    )

    q_num = 1
    for section_title, questions in SECTIONS:
        doc.add_heading(section_title, level=1)
        for question, answer in questions:
            add_answer(doc, q_num, question, answer)
            q_num += 1

    doc.add_heading("Final Advice For The Defense", level=1)
    for item in [
        "Start with the problem, then explain the solution.",
        "When asked about ML, say it supports decisions but does not replace teacher judgment.",
        "Be honest about limitations, especially limited real-world data and local XAMPP deployment.",
        "Know your three users: student, teacher, and administrator.",
        "Be ready to demonstrate login, quiz, progress, teacher analytics, reports, and admin management.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("EduTrack Ghana Defense Prep")
    set_run(fr, 9, False, "555555")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT.resolve())
